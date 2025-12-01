# backend/services/score_resume/extract.py
import io
import re
import openpyxl
import pdfplumber
import hashlib
import asyncio
from dateutil.relativedelta import relativedelta
from datetime import datetime, timedelta
from typing import Optional, Literal
from concurrent.futures import ThreadPoolExecutor
from backend.core.openai_config import get_openai_client

# ============================================
# ✅ LangChain用のインポート
# ============================================
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from pydantic import BaseModel, Field, field_validator, SecretStr

# ============================================
# ✅ GPT呼び出し
# ============================================

client = get_openai_client()

# ============================================
# ✅ LangChain: 構造化出力用のスキーマ定義
# ============================================

class PersonInfo(BaseModel):
    """履歴書から抽出する個人情報"""
    name: Optional[str] = Field(None, description="氏名（フルネーム）。見つからない場合はNone")
    gender: Literal["男性", "女性", "その他"] = Field("その他", description="性別。必ず「男性」「女性」「その他」のいずれか")
    
    @field_validator('gender')
    @classmethod
    def normalize_gender(cls, v: str) -> str:
        """性別を正規化"""
        if not v:
            return "その他"
        v_str = str(v).strip()
        if "男" in v_str and "女" not in v_str:
            return "男性"  # ✅ 変更
        elif "女" in v_str:
            return "女性"  # ✅ 変更
        return "その他"  # ✅ 変更

# ============================================
# 🚀 LangChainモデルのセットアップ
# ============================================

# LangChainモデルのセットアップ（APIキーは自動読み込み）
llm = ChatOpenAI(
    model="gpt-3.5-turbo",
    temperature=0
)

# 構造化出力用のLLM
structured_llm = llm.with_structured_output(PersonInfo, method="function_calling")

# ============================================
# 🚀 パフォーマンス最適化: キャッシング
# ============================================

_extract_cache: dict[str, tuple[str, datetime]] = {}
CACHE_TTL_SECONDS = 3600  # 1時間

def _cache_key(text: str, operation: str) -> str:
    """テキストと操作からキャッシュキーを生成"""
    content = f"{operation}:{text[:500]}"  # 最初の500文字でキャッシュ
    return hashlib.md5(content.encode('utf-8')).hexdigest()

def _get_cached(key: str) -> Optional[str]:
    """キャッシュから取得"""
    if key in _extract_cache:
        result, timestamp = _extract_cache[key]
        if datetime.now() - timestamp < timedelta(seconds=CACHE_TTL_SECONDS):
            print(f"✅ キャッシュヒット: {key[:8]}...")
            return result
        else:
            del _extract_cache[key]
    return None

def _set_cache(key: str, value: str):
    """キャッシュに保存"""
    _extract_cache[key] = (value, datetime.now())

# ============================================
# 🚀 パフォーマンス最適化: 並列実行
# ============================================

_executor = ThreadPoolExecutor(max_workers=3)

async def _call_gpt_async(prompt: str, model: str = "gpt-3.5-turbo", temperature: float = 0.2) -> str:
    """GPTを非同期で呼び出す"""
    # キャッシュチェック
    cache_k = _cache_key(prompt, model)
    cached = _get_cached(cache_k)
    if cached:
        return cached
    
    # 非同期実行
    loop = asyncio.get_event_loop()
    response = await loop.run_in_executor(
        _executor,
        lambda: client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature,
        )
    )
    
    content = response.choices[0].message.content
    result = content.strip() if content else ""
    _set_cache(cache_k, result)
    return result

# ============================================
# 🧠 履歴書からテキストの抽出
# ============================================

def extract_resume_text_from_pdf(file_stream: io.BytesIO) -> str:
    try:
        with pdfplumber.open(file_stream) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        return normalize_pdf_text(text)
    except Exception as e:
        print(f"❌ PDF抽出エラー: {e}")
        return ""

def extract_resume_text_from_docx(file_stream):
    from docx import Document

    doc = Document(file_stream)
    lines = []

    # ① 段落（paragraph）を抽出
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # ② 表（table）を抽出
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                lines.append(" ".join(row_text))

    return "\n".join(lines)

def extract_resume_text_from_xlsx(file_stream: io.BytesIO) -> str:
    try:
        wb = openpyxl.load_workbook(file_stream, data_only=True, read_only=True)  # 🚀 read_only=True で高速化
        text = ""

        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None:
                        text += str(cell).strip() + "\n"

        return text
    except Exception as e:
        print(f"❌ XLSX抽出エラー: {e}")
        return ""
    
def normalize_pdf_text(text: str) -> str:
    """
    🚀 最適化: 正規表現を事前コンパイル
    """
    text = text.replace('\u3000', ' ')
    text = re.sub(r'(?<=[^\n])\n(?=[^\n])', '', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text.strip()

import re
from datetime import datetime

def extract_birth_date(text: str) -> str | None:
    """
    履歴書から生年月日を抽出
    対応形式:
    - 1990年1月1日
    - 1990/01/01
    - 1990-01-01
    - 平成2年1月1日
    """
    
    # 西暦形式
    patterns = [
        r'(\d{4})年(\d{1,2})月(\d{1,2})日',
        r'(\d{4})/(\d{1,2})/(\d{1,2})',
        r'(\d{4})-(\d{1,2})-(\d{1,2})',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            year, month, day = match.groups()
            try:
                date_obj = datetime(int(year), int(month), int(day))
                return date_obj.strftime('%Y-%m-%d')
            except ValueError:
                continue
    
    # 和暦対応（平成・令和）
    wareki_pattern = r'(平成|令和)(\d{1,2})年(\d{1,2})月(\d{1,2})日'
    match = re.search(wareki_pattern, text)
    if match:
        era, era_year, month, day = match.groups()
        
        # 平成 → 1988年基準、令和 → 2018年基準
        base_year = 1988 if era == '平成' else 2018
        year = base_year + int(era_year)
        
        try:
            date_obj = datetime(year, int(month), int(day))
            return date_obj.strftime('%Y-%m-%d')
        except ValueError:
            pass
    
    return None

# ============================================
# 🧠 LangChainを使った名前・性別の抽出（高速化版）
# ============================================

# プロンプトテンプレート
person_info_prompt = ChatPromptTemplate.from_messages([
    ("system", "あなたは履歴書から個人情報を正確に抽出するAIです。"),
    ("user", """以下の履歴書から氏名と性別を抽出してください。

履歴書:
{text}

抽出ルール:
- 氏名: フルネームで返す（例: 山田太郎）。見つからない場合はnull
- 性別: 必ず「男性」「女性」「その他」のいずれか1つのみ
  * 「性別：男」「性別：男性」「男性」などは全て「男性」
  * 「性別：女」「性別：女性」「女性」などは全て「女性」
  * 性別の記載がない場合のみ「その他」
""")
])

# LCELチェーン構築
person_info_chain = person_info_prompt | structured_llm

async def extract_person_info_async(text: str) -> tuple[Optional[str], str]:
    """
    🚀 LangChainで名前と性別を並列抽出（高速化）
    
    Returns:
        (name, gender): 名前と性別のタプル
    """
    # キャッシュチェック
    cache_k = _cache_key(text[:2000], "person_info")
    cached = _get_cached(cache_k)
    if cached:
        parts = cached.split("|")
        return (parts[0] if parts[0] != "None" else None, parts[1] if len(parts) > 1 else "不明")
    
    try:
        # LangChainチェーンを非同期実行
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            _executor,
            lambda: person_info_chain.invoke({"text": text[:2000]})
        )
        
        # Ensure result is a PersonInfo instance
        if isinstance(result, dict):
            result = PersonInfo(**result)
        
        name = result.name
        gender = result.gender  # バリデーターで正規化済み
        
        # キャッシュに保存
        _set_cache(cache_k, f"{name}|{gender}")
        
        print(f"✅ LangChain抽出成功: 氏名={name}, 性別={gender}")
        return (name, gender)
        
    except Exception as e:
        print(f"❌ LangChain抽出エラー: {e}")
        return (None, "不明")

def extract_person_info(text: str) -> tuple[Optional[str], str]:
    """
    同期版ラッパー（イベントループ対応）
    
    使用例:
    name, gender = extract_person_info(resume_text)
    """
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(asyncio.run, extract_person_info_async(text))
                return future.result()
        else:
            return loop.run_until_complete(extract_person_info_async(text))
    except RuntimeError:
        return asyncio.run(extract_person_info_async(text))

# ============================================
# 🧠 履歴書から志望動機の抽出（最適化版）
# ============================================

async def _extract_motivation_async(text: str) -> str:
    """
    🚀 最適化: 非同期版 + キャッシング + プロンプト短縮
    """
    if not text or not text.strip():
        return ""

    # 🚀 プロンプトを簡潔化（トークン削減）
    prompt = f"""
以下の履歴書から「志望動機」または「自己PR」の部分のみを抽出してください。
見つからない場合は空文字を返してください。

履歴書:
{text[:2000]}

抽出結果:
"""

    try:
        return await _call_gpt_async(prompt, model="gpt-3.5-turbo", temperature=0.2)
    except Exception as e:
        print(f"❌ 志望動機抽出に失敗: {e}")
        return ""

def extract_motivation(text: str) -> str:
    """同期版ラッパー（イベントループ対応）"""
    try:
        # 既存のイベントループがある場合
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # ループが実行中の場合は直接awaitできないので、新しいスレッドで実行
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(asyncio.run, _extract_motivation_async(text))
                return future.result()
        else:
            return loop.run_until_complete(_extract_motivation_async(text))
    except RuntimeError:
        # イベントループが存在しない場合
        return asyncio.run(_extract_motivation_async(text))

async def _summarize_motivation_async(text: str, max_length: int = 100) -> str:
    """
    🚀 最適化: 非同期版 + キャッシング + プロンプト短縮
    """
    # 🚀 プロンプトを簡潔化
    prompt = f"""
以下の志望動機を{max_length}文字以内で要約してください。

志望動機:
{text}

要約（{max_length}文字以内）:
"""

    return await _call_gpt_async(prompt, model="gpt-3.5-turbo", temperature=0.3)

def summarize_motivation(text: str, max_length: int = 100) -> str:
    """同期版ラッパー（イベントループ対応）"""
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(asyncio.run, _summarize_motivation_async(text, max_length))
                return future.result()
        else:
            return loop.run_until_complete(_summarize_motivation_async(text, max_length))
    except RuntimeError:
        return asyncio.run(_summarize_motivation_async(text, max_length))

# ============================================
# 🧠 履歴書から職務経歴の抽出（最適化版）
# ============================================

async def _extract_work_experience_async(text: str) -> str:
    """
    🚀 最適化: 非同期版 + キャッシング + プロンプト短縮
    """
    if not text or not text.strip():
        return ""

    # 🚀 プロンプトを簡潔化
    prompt = f"""
以下の履歴書から「職務経歴」「業務内容」「経歴」に該当する部分を抽出してください。
見つからない場合は空文字を返してください。

履歴書:
{text[:2000]}

抽出結果:
"""

    try:
        return await _call_gpt_async(prompt, model="gpt-3.5-turbo", temperature=0.2)
    except Exception as e:
        print(f"❌ 職務経歴抽出に失敗: {e}")
        return ""

def extract_work_experience(text: str) -> str:
    """同期版ラッパー（イベントループ対応）"""
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(asyncio.run, _extract_work_experience_async(text))
                return future.result()
        else:
            return loop.run_until_complete(_extract_work_experience_async(text))
    except RuntimeError:
        return asyncio.run(_extract_work_experience_async(text))
    
async def _summarize_work_experience_async(text: str, max_length: int = 150) -> str:
    """
    🚀 最適化: 非同期版 + キャッシング + プロンプト短縮
    """
    # 🚀 プロンプトを簡潔化
    prompt = f"""
以下の職務経歴を{max_length}文字以内で要約してください。

職務経歴:
{text}

要約（{max_length}文字以内）:
"""

    return await _call_gpt_async(prompt, model="gpt-3.5-turbo", temperature=0.3)

def summarize_work_experience(text: str, max_length: int = 150) -> str:
    """同期版ラッパー（イベントループ対応）"""
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(asyncio.run, _summarize_work_experience_async(text, max_length))
                return future.result()
        else:
            return loop.run_until_complete(_summarize_work_experience_async(text, max_length))
    except RuntimeError:
        return asyncio.run(_summarize_work_experience_async(text, max_length))

# ============================================
# 🚀 新機能: 志望動機と職務経歴を並列抽出
# ============================================

async def extract_motivation_and_experience_async(text: str) -> tuple[str, str]:
    """
    🚀 並列実行: 志望動機と職務経歴を同時に抽出
    処理時間を約半分に短縮
    """
    motivation_task = _extract_motivation_async(text)
    experience_task = _extract_work_experience_async(text)
    
    motivation, experience = await asyncio.gather(
        motivation_task,
        experience_task
    )
    
    return motivation, experience

def extract_motivation_and_experience(text: str) -> tuple[str, str]:
    """
    同期版ラッパー（イベントループ対応）
    
    使用例:
    motivation, experience = extract_motivation_and_experience(resume_text)
    """
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(asyncio.run, extract_motivation_and_experience_async(text))
                return future.result()
        else:
            return loop.run_until_complete(extract_motivation_and_experience_async(text))
    except RuntimeError:
        return asyncio.run(extract_motivation_and_experience_async(text))

# ============================================
# 🧠 履歴書から社会人歴の抽出
# ============================================

# 正規表現の事前コンパイル
PATTERN_YM_JP     = re.compile(r"(\d{4})年\s*(\d{1,2})月")
PATTERN_YM_DASH   = re.compile(r"(\d{4})-(\d{1,2})")
PATTERN_YM_SLASH  = re.compile(r"(\d{4})/(\d{1,2})")
PATTERN_YM_DOT    = re.compile(r"(\d{4})\.(\d{1,2})")
PATTERN_YM_SPACE  = re.compile(r"(\d{4})\s+(\d{1,2})")

def parse_date(date_str):
    if not date_str:
        return None

    s = date_str.strip()

    # 軽い正規化（全角 → 半角・余計な文字除去）
    s = s.replace("（", "(").replace("）", ")")
    s = re.sub(r"[^\d年月Present今現在至今/.\-\s]", "", s)
    s = re.sub(r"\s+", " ", s)

    # 現在扱い
    if s in ["現在", "今", "現職", "至今", "Present", "present", "PRESENT"]:
        return datetime.today()

    # YYYY年MM月
    m = PATTERN_YM_JP.match(s)
    if m:
        return datetime(int(m.group(1)), int(m.group(2)), 1)

    # YYYY-MM
    m = PATTERN_YM_DASH.match(s)
    if m:
        return datetime(int(m.group(1)), int(m.group(2)), 1)

    # YYYY/MM
    m = PATTERN_YM_SLASH.match(s)
    if m:
        return datetime(int(m.group(1)), int(m.group(2)), 1)

    # YYYY.MM
    m = PATTERN_YM_DOT.match(s)
    if m:
        return datetime(int(m.group(1)), int(m.group(2)), 1)

    # YYYY MM
    m = PATTERN_YM_SPACE.match(s)
    if m:
        return datetime(int(m.group(1)), int(m.group(2)), 1)

    return None

def calculate_total_experience(work_histories):
    """
    社会人歴の計算（最適化済み）
    """
    periods = []

    for history in work_histories:
        start = parse_date(history.start_date)
        end = parse_date(history.end_date) or datetime.today()

        if start and end:
            periods.append((start, end))

    # 重複期間のマージ
    periods.sort()
    merged = []

    for start, end in periods:
        if not merged:
            merged.append((start, end))
        else:
            last_start, last_end = merged[-1]
            if start <= last_end:
                merged[-1] = (last_start, max(last_end, end))
            else:
                merged.append((start, end))

    # 総経験年数を月単位で計算
    total_months = sum((relativedelta(end, start).years * 12 + relativedelta(end, start).months for start, end in merged))

    return round(total_months / 12, 1)

# ============================================
# 🚀 新機能: 全抽出処理を並列実行（LangChain版）
# ============================================

async def extract_all_resume_info_async(text: str) -> dict:
    """
    🚀 すべての抽出処理を並列実行（LangChain使用）
    
    Returns:
        {
            "name": str,
            "gender": str,
            "motivation": str,
            "work_experience": str
        }
    """
    # 並列実行するタスクを準備
    person_info_task = extract_person_info_async(text)
    motivation_task = _extract_motivation_async(text)
    experience_task = _extract_work_experience_async(text)
    
    # すべてを並列実行
    results = await asyncio.gather(
        person_info_task,
        motivation_task,
        experience_task,
        return_exceptions=True
    )
    
    # エラーハンドリング
    person_info_result = results[0]
    if isinstance(person_info_result, Exception):
        name, gender = None, "不明"
    else:
        if isinstance(person_info_result, BaseException):
            name, gender = None, "不明"
        else:
            name, gender = person_info_result
    
    motivation = results[1] if not isinstance(results[1], Exception) else ""
    experience = results[2] if not isinstance(results[2], Exception) else ""
    
    return {
        "name": name,
        "gender": gender,
        "motivation": motivation,
        "work_experience": experience,
    }

def extract_all_resume_info(text: str) -> dict:
    """
    同期版ラッパー（イベントループ対応）
    
    使用例:
    info = extract_all_resume_info(resume_text)
    print(info["name"], info["motivation"])
    """
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(asyncio.run, extract_all_resume_info_async(text))
                return future.result()
        else:
            return loop.run_until_complete(extract_all_resume_info_async(text))
    except RuntimeError:
        return asyncio.run(extract_all_resume_info_async(text))